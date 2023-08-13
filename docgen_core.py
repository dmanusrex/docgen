# DocGen - https://github.com/dmanusrex/docgen
#
# Copyright (C) 2023 - Darren Richer
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND,
# EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF
# MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT.
# IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
# DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
# OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE
# OR OTHER DEALINGS IN THE SOFTWARE.

''' Support functions for the main application '''

import pandas as pd
from threading import Thread
import tkinter as tk
import numpy as np
import os
from slugify import slugify

from datetime import datetime
from typing import List
from config import docgenConfig
from docx import Document
import docx
from docxcompose.composer import Composer

import logging

class docgenCore:

    def __init__(self, club: str, club_data_set : pd.DataFrame, config: docgenConfig, **kwargs):
        self._club_data_full = club_data_set
        self._club_data = self._club_data_full.query("Current_CertificationLevel not in ['LEVEL IV - GREEN PIN','LEVEL V - BLUE PIN']")
        self.club_code = club
      
        self._config = config


    def _is_valid_date(self, date_string) -> bool:
        if pd.isnull(date_string): return False
        if date_string == "0001-01-01": return False 
        try:
            datetime.strptime(date_string, '%Y-%m-%d')
            return True
        except ValueError:
            return False
    
    def _get_date(self, date_string) -> str: 
        if pd.isnull(date_string): return ""
        if date_string == "0001-01-01": return "" 
        return date_string
  
    def _count_signoffs(self, clinic_date_1, clinic_date_2) -> int:
        count = 0
        if self._is_valid_date(clinic_date_1): count += 1
        if self._is_valid_date(clinic_date_2): count += 1
        return count
    
    def add_clinic(self, table, clinic_name, clinic_date, signoff_1, signoff_2) -> None:
        row = table.add_row().cells
        row[0].text = clinic_name
        row[1].text = self._get_date(clinic_date)
        row[2].text = self._get_date(signoff_1)
        row[3].text = self._get_date(signoff_2)
        row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
  
    def dump_data_docx(self, club_fullname: str, reportdate: str) -> List:
        '''Produce the Word Document for the club and return a list of files'''
 
        _report_directory = self._config.get_str("report_directory")
        _email_list_csv = self._config.get_str("email_list_csv")
        csv_list = []    # CSV entries for email list (Lastname, Firstname, E-Mail address and Filename)
 
        for index, entry in self._club_data.iterrows():

            # create a filename from the last and firstnames using slugify and the report directory

            filename = os.path.abspath(os.path.join(_report_directory, slugify(entry["Last Name"] + "_" + entry["First Name"]) + ".docx"))
            csv_list.append([entry["Last Name"], entry["First Name"], entry["Email"], filename])

            doc = Document()

            doc.add_heading("2023/24 Officials Development", 0)
 
            p = doc.add_paragraph()
            p.add_run("Report Date: "+reportdate)
            p.add_run("\n\nName: "+ entry["Last Name"] + ", " + entry["First Name"] + " (SNC ID # " + entry["Registration Id"] + ")")
            p.add_run("\n\nClub: "+ club_fullname + " (" + self.club_code + ")")
            p.add_run("\n\nCurrent Certification Level: ")
            p.add_run("NONE" if pd.isnull(entry["Current_CertificationLevel"]) else entry["Current_CertificationLevel"])



            table = doc.add_table(rows=1, cols=4)
            row = table.rows[0].cells
            row[0].text = "Clinic"
            row[1].text = "Clinic Date"
            row[2].text = "Sign Off #1"
            row[3].text = "Sign Off #2" 
            row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            row[1].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            row[3].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER


            self.add_clinic(table, "Intro to Swimming", entry["Introduction to Swimming Officiating-ClinicDate"], entry["Introduction to Swimming Officiating-Deck Evaluation #1 Date"], entry["Introduction to Swimming Officiating-Deck Evaluation #2 Date"])
            self.add_clinic(table, "Safety Marshal", entry["Safety Marshal-ClinicDate"], "N/A", "N/A")
            self.add_clinic(table, "Stroke & Turn", entry["Judge of Stroke/Inspector of Turns-ClinicDate"], entry["Judge of Stroke/Inspector of Turns-Deck Evaluation #1 Date"], entry["Judge of Stroke/Inspector of Turns-Deck Evaluation #2 Date"])
            self.add_clinic(table, "Chief Timekeeper", entry["Chief Timekeeper-ClinicDate"], entry["Chief Timekeeper-Deck Evaluation #1 Date"], entry["Chief Timekeeper-Deck Evaluation #2 Date"])
            self.add_clinic(table, "Admin Desk (Clerk)", entry["Clerk of Course-ClinicDate"], entry["Clerk of Course-Deck Evaluation #1 Date"], entry["Clerk of Course-Deck Evaluation #2 Date"])
            self.add_clinic(table, "Meet Manager", entry["Meet Manager-ClinicDate"], entry["Meet Manager-Deck Evaluation #1 Date"], entry["Meet Manager-Deck Evaluation #2 Date"])
            self.add_clinic(table, "Starter", entry["Starter-ClinicDate"], entry["Starter-Deck Evaluation #1 Date"], entry["Starter-Deck Evaluation #2 Date"])
            self.add_clinic(table, "CFJ/CJE", entry["Chief Finish Judge/Chief Judge-ClinicDate"], entry["Chief Finish Judge/Chief Judge-Deck Evaluation #1 Date"], entry["Chief Finish Judge/Chief Judge-Deck Evaluation #2 Date"])
            self.add_clinic(table, "Chief Recorder/Recorder", entry["Recorder-Scorer-ClinicDate"], "N/A", "N/A")
            self.add_clinic(table, "Referee", entry["Referee-ClinicDate"], "N/A", "N/A")
            self.add_clinic(table, "Para eModule", entry["Para Swimming eModule-ClinicDate"], "N/A", "N/A")
            
            table.style = "Light Grid Accent 5"
            table.autofit = True

            # Add logic to define pathway progression

            doc.add_heading("Recommended Actions", 2)

            # For NoLevel officials, add a section to identify what they need to do to get to Level I

            if pd.isnull(entry["Current_CertificationLevel"]):
                if entry["Introduction to Swimming Officiating"].lower() == "no":
                    doc.add_paragraph("Take Introduction to Swimming Officiating Clinic and obtain sign-offs", style="List Bullet")
                else:
                    # if < 2 clinics tell user to get remaining sign-offs (2 - # of clinics)
                    num_signoffs = self._count_signoffs(entry["Introduction to Swimming Officiating-Deck Evaluation #1 Date"],entry["Introduction to Swimming Officiating-Deck Evaluation #2 Date"])
                    if num_signoffs < 2:
                        doc.add_paragraph(f"Obtain {2-num_signoffs} sign-off(s) for Introduction to Swimming Officiating", style="List Bullet")
                   
                if entry["Safety Marshal"].lower() == "no":
                        doc.add_paragraph("Take Safety Marshal Clinc", style="List Bullet")

            # For Level I officials - check if they have stroke & turn and have completed 2 sign-offs

            if entry["Current_CertificationLevel"] == "LEVEL I - RED PIN":
                intro_signoffs = self._count_signoffs(entry["Introduction to Swimming Officiating-Deck Evaluation #1 Date"],entry["Introduction to Swimming Officiating-Deck Evaluation #2 Date"])
                if intro_signoffs < 2:
                    doc.add_paragraph(f"Obtain {2-intro_signoffs} sign-off(s) for Introduction to Swimming Officiating", style="List Bullet")

                if entry["Judge of Stroke/Inspector of Turns"].lower() == "no":
                    doc.add_paragraph("Take Judge of Stroke/Inspector of Turns Clinic", style="List Bullet")
                else:
                    num_signoffs = self._count_signoffs(entry["Judge of Stroke/Inspector of Turns-Deck Evaluation #1 Date"],entry["Judge of Stroke/Inspector of Turns-Deck Evaluation #2 Date"])
                    if num_signoffs < 2:
                        doc.add_paragraph(f"Obtain {2-num_signoffs} sign-off(s) for Judge of Stroke/Inspector of Turns", style="List Bullet")
                    # At this point we know they have stroke & turn and intro.  Determine Level II recommendations.
                    if intro_signoffs + num_signoffs >= 3:   # They have completed or nearly completed "core" requirements
                        # Check if they have any other clinics
                        if entry["Chief Timekeeper"].lower() == "no" and entry["Clerk of Course"].lower() == "no" and entry["Meet Manager"].lower() == "no" and entry["Starter"].lower() == "no" and entry["Chief Finish Judge/Chief Judge"].lower() == "no": 
                            doc.add_paragraph("Take a Level II clinic (CT, MM, CFJ/CJE, Admin Desk or Starter) and obtain sign-offs", style="List Bullet")
                        else:
                            doc.add_paragraph("Obtain sign-offs on at least 1 Level II clinic (CT, MM, CFJ/CJE, Admin Desk or Starter)", style="List Bullet")
            try:
                doc.save(filename)

            except Exception as e:
                logging.info(f'Error processing offiical {entry["Last Name"]}, {entry["First Name"]}: {type(e).__name__} - {e}')

        return csv_list

class TextHandler(logging.Handler):
    # This class allows you to log to a Tkinter Text or ScrolledText widget
    # Adapted from Moshe Kaplan: https://gist.github.com/moshekaplan/c425f861de7bbf28ef06

    def __init__(self, text):
        # run the regular Handler __init__
        logging.Handler.__init__(self)
        # Store a reference to the Text it will log to
        self.text = text

    def emit(self, record):
        msg = self.format(record)
        def append():
            self.text.configure(state='normal')
            self.text.insert(tk.END, msg + '\n')
            self.text.configure(state='disabled')
            # Autoscroll to the bottom
            self.text.yview(tk.END)
        # This is necessary because we can't modify the Text from other threads
        self.text.after(0, append)

class Data_Loader(Thread):
    '''Load Data files'''
    def __init__(self, config: docgenConfig):
        super().__init__()
        self._config = config
        self.df : pd.DataFrame 
        self.affiliates : pd.DataFrame 

    def run(self):
        html_file = self._config.get_str("officials_list")
        self.club_list_names_df = pd.DataFrame
        self.club_list_names = []
        logging.info("Loading RTR Data")
        try:
            self.df = pd.read_html(html_file)[0]
        except:
            logging.info("Unable to load data file")
            self.df = pd.DataFrame
            self.affiliates = pd.DataFrame
            return
        self.df.columns = self.df.iloc[0]   # The first row is the column names
        self.df = self.df[1:]

        # Club Level exports include blank rows, purge those out
        self.df.drop(index=self.df[self.df['Registration Id'].isnull()].index, inplace=True)

        # The RTR has 2 types of "empty" dates.  One is blank the other is 0001-01-01.  Fix that.
        self.df.replace('0001-01-01', np.nan, inplace=True)

        # The RTR export is inconsistent on column values for certifications. Fix that.
        self.df.replace('Yes','yes', inplace=True)    # We don't use the no value so no need to fix it 

        logging.info("Loaded %d officials" % self.df.shape[0])

        logging.info("Loading Complete")

class Generate_Reports(Thread):
    def __init__(self, df: pd.DataFrame, config: docgenConfig):
        super().__init__()
        self._df : pd.DataFrame = df
        self._config : docgenConfig = config

    def run(self):
        logging.info("Reporting in Progress...")

        _report_directory = self._config.get_str("report_directory")
        _report_file_docx = self._config.get_str("report_file_docx")
        _full_report_file = os.path.abspath(os.path.join(_report_directory, _report_file_docx))
        _email_list_csv = self._config.get_str("email_list_csv")
        _full_csv_file = os.path.abspath(os.path.join(_report_directory, _email_list_csv))


        club_list_names_df = self._df.loc[self._df['AffiliatedClubs'].isnull(),['ClubCode','Club']].drop_duplicates()
        club_list_names = club_list_names_df.values.tolist()
        club_list_names.sort(key=lambda x:x[0])

        club_summaries = []

        status_values = ["Active"]
        if self._config.get_bool("incl_inv_pending"):
            status_values.append("Invoice Pending")
        if self._config.get_bool("incl_account_pending"):
            status_values.append("Account Pending")
        if self._config.get_bool("incl_pso_pending"):
            status_values.append("PSO Pending")

        report_time = datetime.now().strftime("%B %d %Y %I:%M%p")

        all_csv_entries = []

        for club, club_full in club_list_names:
            logging.info("Processing %s" % club_full)
            club_data = self._df[(self._df["ClubCode"] == club)]
            club_data = club_data[club_data["Status"].isin(status_values)]
            club_stat = docgenCore(club, club_data, self._config)
            club_csv = club_stat.dump_data_docx(club_full, report_time)
            all_csv_entries.extend(club_csv)
            club_summaries.append ([club, club_full, club_stat])

        # Create the email list CSV file    
        # 
        # The email list is a CSV file with the following columns:
        #  Last Name, First Name, E-Mail address, Filename

        logging.info("Creating email list CSV file")
        email_list_df = pd.DataFrame(all_csv_entries, columns=["Last Name", "First Name", "EMail", "Filename"])

        try:
            email_list_df.to_csv(_full_csv_file, index=False)
        except Exception as e:
            logging.info("Unable to save email list: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))
        
        # Create the master document

        logging.info("Creating master document")

        number_of_sections=len(all_csv_entries)
        master = Document()
        composer = Composer(master)
        for i in range(0, number_of_sections):
            doc_temp = Document(all_csv_entries[i][3])
            doc_temp.add_page_break()
            composer.append(doc_temp)

        try:
            composer.save(_full_report_file)
        except Exception as e:
            logging.info("Unable to save full report: {}".format(type(e).__name__))
            logging.info("Exception message: {}".format(e))

        logging.info("Report Complete")


